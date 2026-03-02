import io
from pathlib import Path
from flask import Blueprint, request, jsonify, g, send_file
from api.auth import login_required
from core.supabase_client import supabase
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

export_bp = Blueprint('export', __name__)


@export_bp.route('/<int:id>/export', methods=['GET'])
@login_required
def export_resume(id):
    """
    E-01 匯出履歷文件 — 從 resume_optimization 取最新版本
    支援 format=pdf (預設) | docx | json
    可帶 ?version=2.0 指定版本
    """
    try:
        user_id = g.db_user_id
        if user_id is None:
            return jsonify({'error': 'User not found in DB'}), 403
        fmt = request.args.get('format', 'pdf').lower()
        target_version = request.args.get('version')  # optional

        if fmt not in ('pdf', 'docx', 'json'):
            return jsonify({'error': f'Unsupported format: {fmt}. Use pdf, docx, or json'}), 400

        # 從 resume_optimization 取資料
        query = (
            supabase.table("resume_optimization")
            .select("*")
            .eq("resume_id", id)
            .eq("user_id", user_id)
        )

        if target_version:
            query = query.eq("optimization_version", target_version).single()
        else:
            # 取最新版本
            query = query.order("optimization_version", desc=True).limit(1)

        response = query.execute()

        opt_data = response.data
        if isinstance(opt_data, list):
            opt_data = opt_data[0] if opt_data else None
        if not opt_data:
            return jsonify({'error': 'No optimized resume found for this resume_id'}), 404

        # 將 resume_optimization 的個別欄位組裝成 structured_data 格式
        # 讓下游 PDF / DOCX builder 可以沿用
        resume_data = _build_structured_from_optimization(opt_data)

        # --- JSON ---
        if fmt == 'json':
            return jsonify({
                'resume_id': id,
                'optimization_version': opt_data.get('optimization_version'),
                'format': 'json',
                'data': opt_data
            }), 200

        # --- PDF ---
        if fmt == 'pdf':
            buf = _build_resume_pdf(resume_data)
            return send_file(buf, mimetype='application/pdf',
                             as_attachment=True, download_name=f"resume_{id}.pdf")

        # --- DOCX ---
        buf = _build_resume_docx(resume_data)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"resume_{id}.docx"
        )

    except Exception as e:
        if 'Row not found' in str(e) or '0 rows' in str(e):
            return jsonify({'error': 'No optimized resume found'}), 404
        return jsonify({'error': str(e)}), 500


def _build_structured_from_optimization(opt: dict) -> dict:
    """將 resume_optimization 各欄位組裝成 builder 預期的格式"""
    sd = {}
    if opt.get('professional_summary'):
        sd['summary'] = opt['professional_summary']
    if opt.get('professional_experience'):
        sd['work_experience'] = opt['professional_experience']
    if opt.get('core_skills'):
        sd['skills'] = opt['core_skills']
    if opt.get('projects'):
        sd['projects'] = opt['projects']
    if opt.get('education'):
        sd['education'] = opt['education']
    if opt.get('autobiography'):
        sd['autobiography'] = opt['autobiography']
    return {'structured_data': sd, 'user_id': opt.get('user_id')}


# ──────────────────────────────────────────────
#  PDF Builder (fpdf2 + 系統 CJK 字型)
# ──────────────────────────────────────────────

def _build_resume_pdf(resume_data: dict) -> io.BytesIO:
    sd = resume_data.get('structured_data', {})

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    _use_cjk = False
    _cjk_candidates = [
        "C:/Windows/Fonts/msjh.ttc",
        "C:/Windows/Fonts/mingliu.ttc",
        "C:/Windows/Fonts/simsun.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/noto-cjk/NotoSansCJKtc-Regular.otf",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    ]
    for font_path in _cjk_candidates:
        if Path(font_path).exists():
            try:
                pdf.add_font("cjk", style="", fname=font_path)
                pdf.add_font("cjk", style="B", fname=font_path)
                pdf.add_font("cjk", style="I", fname=font_path)
                pdf.set_font("cjk", size=12)
                _use_cjk = True
                break
            except Exception:
                continue

    if not _use_cjk:
        pdf.set_font("Helvetica", size=12)

    def _set(style="", size=12):
        pdf.set_font("cjk" if _use_cjk else "Helvetica", style=style, size=size)

    def _section_title(title: str):
        _set("B", 14)
        pdf.set_text_color(30, 80, 160)
        pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(30, 80, 160)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(4)
        pdf.set_text_color(0, 0, 0)

    personal = sd.get('personal_info', {})
    name = personal.get('name', resume_data.get('user_id', ''))
    email = personal.get('email', '')
    phone = personal.get('phone', '')
    location = personal.get('location', '')

    _set("B", 22)
    pdf.cell(0, 12, str(name), new_x="LMARGIN", new_y="NEXT", align="C")
    _set("", 10)
    contact_parts = [p for p in [email, phone, location] if p]
    if contact_parts:
        pdf.cell(0, 7, " | ".join(contact_parts), new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    summary = sd.get('summary', '')
    if summary:
        _section_title("Summary")
        _set("", 11)
        pdf.multi_cell(0, 6, str(summary))
        pdf.ln(4)

    experiences = sd.get('work_experience', sd.get('experience', []))
    if experiences:
        _section_title("Work Experience")
        for exp in experiences:
            if isinstance(exp, str):
                # model_output.md 格式：每筆是完整字串
                _set("", 11)
                pdf.multi_cell(0, 6, exp)
                pdf.ln(3)
            elif isinstance(exp, dict):
                _set("B", 11)
                title = exp.get('title', exp.get('position', ''))
                company = exp.get('company', exp.get('organization', ''))
                period = exp.get('period', exp.get('duration', ''))
                pdf.cell(0, 7, f"{title}  —  {company}", new_x="LMARGIN", new_y="NEXT")
                if period:
                    _set("I", 9)
                    pdf.cell(0, 5, str(period), new_x="LMARGIN", new_y="NEXT")
                desc = exp.get('description', exp.get('details', ''))
                if desc:
                    _set("", 10)
                    pdf.multi_cell(0, 5, str(desc))
                pdf.ln(3)

    educations = sd.get('education', [])
    if educations:
        _section_title("Education")
        for edu in educations:
            if isinstance(edu, str):
                _set("", 11)
                pdf.multi_cell(0, 6, edu)
                pdf.ln(3)
            elif isinstance(edu, dict):
                _set("B", 11)
                school = edu.get('school', edu.get('institution', ''))
                degree = edu.get('degree', '')
                major = edu.get('major', edu.get('field', ''))
                period = edu.get('period', edu.get('duration', ''))
                line = school
                if degree:
                    line += f"  —  {degree}"
                if major:
                    line += f" ({major})"
                pdf.cell(0, 7, line, new_x="LMARGIN", new_y="NEXT")
                if period:
                    _set("I", 9)
                    pdf.cell(0, 5, str(period), new_x="LMARGIN", new_y="NEXT")
                details = edu.get('details', '')
                if details:
                    _set("", 10)
                    pdf.multi_cell(0, 5, str(details))
                pdf.ln(3)

    skills = sd.get('skills', [])
    if skills:
        _section_title("Skills")
        _set("", 11)
        pdf.multi_cell(0, 6, " • ".join(str(s) for s in skills) if isinstance(skills, list) else str(skills))
        pdf.ln(4)

    projects = sd.get('projects', [])
    if projects:
        _section_title("Projects")
        for proj in projects:
            if isinstance(proj, str):
                _set("", 11)
                pdf.multi_cell(0, 6, proj)
                pdf.ln(3)
            elif isinstance(proj, dict):
                _set("B", 11)
                pdf.cell(0, 7, str(proj.get('name', '')), new_x="LMARGIN", new_y="NEXT")
                desc = proj.get('description', '')
                if desc:
                    _set("", 10)
                    pdf.multi_cell(0, 5, str(desc))
                pdf.ln(3)

    certs = sd.get('certifications', sd.get('certificates', []))
    if certs:
        _section_title("Certifications")
        _set("", 11)
        for cert in certs:
            label = cert.get('name', '') if isinstance(cert, dict) else cert
            pdf.cell(0, 6, f"• {label}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    buf = io.BytesIO()
    buf.write(pdf.output())
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────
#  DOCX Builder (python-docx)
# ──────────────────────────────────────────────

def _build_resume_docx(resume_data: dict) -> io.BytesIO:
    sd = resume_data.get('structured_data', {})
    doc = Document()

    doc.styles['Normal'].font.size = Pt(11)
    ACCENT = RGBColor(30, 80, 160)

    def _section_heading(title: str):
        h = doc.add_heading(level=2)
        run = h.runs[0] if h.runs else h.add_run(title)
        run.text = title
        run.font.color.rgb = ACCENT
        run.font.size = Pt(14)

    personal = sd.get('personal_info', {})
    name = personal.get('name', str(resume_data.get('user_id', '')))
    email = personal.get('email', '')
    phone = personal.get('phone', '')
    location = personal.get('location', '')

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(name)
    run.bold = True
    run.font.size = Pt(22)

    contact_parts = [p for p in [email, phone, location] if p]
    if contact_parts:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run(' | '.join(contact_parts)).font.size = Pt(10)

    summary = sd.get('summary', '')
    if summary:
        _section_heading('Summary')
        doc.add_paragraph(str(summary))

    experiences = sd.get('work_experience', sd.get('experience', []))
    if experiences:
        _section_heading('Work Experience')
        for exp in experiences:
            if isinstance(exp, str):
                doc.add_paragraph(exp)
            elif isinstance(exp, dict):
                title = exp.get('title', exp.get('position', ''))
                company = exp.get('company', exp.get('organization', ''))
                period = exp.get('period', exp.get('duration', ''))
                p = doc.add_paragraph()
                run = p.add_run(f'{title}  —  {company}')
                run.bold = True
                run.font.size = Pt(11)
                if period:
                    p2 = doc.add_paragraph()
                    run2 = p2.add_run(str(period))
                    run2.italic = True
                    run2.font.size = Pt(9)
                desc = exp.get('description', exp.get('details', ''))
                if desc:
                    doc.add_paragraph(str(desc))

    educations = sd.get('education', [])
    if educations:
        _section_heading('Education')
        for edu in educations:
            if isinstance(edu, str):
                doc.add_paragraph(edu)
            elif isinstance(edu, dict):
                school = edu.get('school', edu.get('institution', ''))
                degree = edu.get('degree', '')
                major = edu.get('major', edu.get('field', ''))
                period = edu.get('period', edu.get('duration', ''))
                line = school
                if degree:
                    line += f'  —  {degree}'
                if major:
                    line += f' ({major})'
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(11)
                if period:
                    p2 = doc.add_paragraph()
                    run2 = p2.add_run(str(period))
                    run2.italic = True
                    run2.font.size = Pt(9)
                details = edu.get('details', '')
                if details:
                    doc.add_paragraph(str(details))

    skills = sd.get('skills', [])
    if skills:
        _section_heading('Skills')
        doc.add_paragraph(' • '.join(str(s) for s in skills) if isinstance(skills, list) else str(skills))

    projects = sd.get('projects', [])
    if projects:
        _section_heading('Projects')
        for proj in projects:
            if isinstance(proj, str):
                doc.add_paragraph(proj)
            elif isinstance(proj, dict):
                p = doc.add_paragraph()
                p.add_run(str(proj.get('name', ''))).bold = True
                desc = proj.get('description', '')
                if desc:
                    doc.add_paragraph(str(desc))

    certs = sd.get('certifications', sd.get('certificates', []))
    if certs:
        _section_heading('Certifications')
        for cert in certs:
            label = cert.get('name', '') if isinstance(cert, dict) else cert
            doc.add_paragraph(f"• {label}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
